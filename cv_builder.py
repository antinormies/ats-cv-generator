from docx import Document
from docx.shared import Pt, Inches, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml.ns import qn, nsdecls
from docx.oxml import OxmlElement, parse_xml
from fpdf import FPDF
import os
import re

from data.personal_info import PERSONAL_INFO


class CVBuilder:
    ATS_KEYWORDS = [
        "Android", "Kotlin", "Java", "Jetpack Compose", "Android SDK",
        "Clean Architecture", "MVVM", "MAD", "Coroutines", "Flow",
        "Hilt", "Dagger", "Koin", "Retrofit", "OkHttp", "REST API",
        "Room", "Firebase", "Firebase Auth", "Firestore", "Crashlytics",
        "CI/CD", "Fastlane", "GitHub Actions", "Google Play",
        "Material Design 3", "Material Design", "JNI", "NDK",
        "llama.cpp", "whisper.cpp", "TFLite", "LiteRT", "ONNX",
        "Machine Learning", "AI", "LLM", "CNN",
        "WorkManager", "Navigation Component", "Paging 3",
        "CameraX", "WebSocket", "Git",
        "Performance Optimization", "Code Review",
        "Agile", "Scrum", "Unit Test", "JUnit", "MockK",
        "SQLite", "DataStore", "StateFlow",
        "Google Play Console", "Play Store",
        "Adaptive Layout", "Jetpack",
        "Deep Learning", "On-Device", "Inference",
    ]

    def __init__(self, info: dict = None):
        self.info = info or PERSONAL_INFO

    def _set_cell_border(self, cell, **kwargs):
        tc = cell._tc
        tcPr = tc.get_or_add_tcPr()
        tcBorders = OxmlElement("w:tcBorders")
        for edge, val in kwargs.items():
            element = OxmlElement(f"w:{edge}")
            for attr, value in val.items():
                element.set(qn(f"w:{attr}"), str(value))
            tcBorders.append(element)
        tcPr.append(tcBorders)

    def _add_hyperlink(self, paragraph, text: str, url: str, font_size=9):
        part = paragraph.part
        r_id = part.relate_to(url, "http://schemas.openxmlformats.org/officeDocument/2006/relationships/hyperlink", is_external=True)
        hyperlink = OxmlElement("w:hyperlink")
        hyperlink.set(qn("r:id"), r_id)
        new_run = OxmlElement("w:r")
        rPr = OxmlElement("w:rPr")
        c = OxmlElement("w:color")
        c.set(qn("w:val"), "2255CC")
        rPr.append(c)
        u = OxmlElement("w:u")
        u.set(qn("w:val"), "single")
        rPr.append(u)
        sz = OxmlElement("w:sz")
        sz.set(qn("w:val"), str(int(font_size * 2)))
        rPr.append(sz)
        new_run.append(rPr)
        t = OxmlElement("w:t")
        t.text = text
        new_run.append(t)
        hyperlink.append(new_run)
        paragraph._p.append(hyperlink)
        return hyperlink

    def _add_hr(self, doc):
        p = doc.add_paragraph()
        p.paragraph_format.space_before = Pt(2)
        p.paragraph_format.space_after = Pt(2)
        pPr = p._p.get_or_add_pPr()
        pBdr = OxmlElement("w:pBdr")
        bottom = OxmlElement("w:bottom")
        bottom.set(qn("w:val"), "single")
        bottom.set(qn("w:sz"), "4")
        bottom.set(qn("w:space"), "1")
        bottom.set(qn("w:color"), "333333")
        pBdr.append(bottom)
        pPr.append(pBdr)

    def build_docx(self, output_path: str) -> str:
        doc = Document()
        section = doc.sections[0]
        section.top_margin = Inches(0.5)
        section.bottom_margin = Inches(0.5)
        section.left_margin = Inches(0.7)
        section.right_margin = Inches(0.7)

        style = doc.styles["Normal"]
        font = style.font
        font.name = "Calibri"
        font.size = Pt(10.5)
        font.color.rgb = RGBColor(0x33, 0x33, 0x33)
        style.paragraph_format.space_after = Pt(3)
        style.paragraph_format.space_before = Pt(0)
        style.paragraph_format.line_spacing = 1.08

        # --- HEADER ---
        name_p = doc.add_paragraph()
        name_p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        name_run = name_p.add_run(self.info["name"])
        name_run.bold = True
        name_run.font.size = Pt(20)
        name_run.font.color.rgb = RGBColor(0x1A, 0x1A, 0x2E)

        title_p = doc.add_paragraph()
        title_p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        title_run = title_p.add_run(self.info["title"])
        title_run.font.size = Pt(11)
        title_run.font.color.rgb = RGBColor(0x55, 0x55, 0x55)

        contact_p = doc.add_paragraph()
        contact_p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        contact_items = []
        if self.info["location"]:
            r = contact_p.add_run(self.info["location"])
            r.font.size = Pt(9)
            r.font.color.rgb = RGBColor(0x66, 0x66, 0x66)
            contact_items.append(True)
        if self.info["linkedin"]:
            if contact_items:
                r = contact_p.add_run("  •  ")
                r.font.size = Pt(9)
                r.font.color.rgb = RGBColor(0x66, 0x66, 0x66)
            self._add_hyperlink(contact_p, "LinkedIn", self.info["linkedin"], font_size=9)
            contact_items.append(True)
        if self.info["github"]:
            if contact_items:
                r = contact_p.add_run("  •  ")
                r.font.size = Pt(9)
                r.font.color.rgb = RGBColor(0x66, 0x66, 0x66)
            self._add_hyperlink(contact_p, "GitHub", self.info["github"], font_size=9)
            contact_items.append(True)
        if self.info["email"]:
            if contact_items:
                r = contact_p.add_run("  •  ")
                r.font.size = Pt(9)
                r.font.color.rgb = RGBColor(0x66, 0x66, 0x66)
            self._add_hyperlink(contact_p, self.info["email"], f"mailto:{self.info['email']}", font_size=9)
            contact_items.append(True)
        if self.info["website"]:
            if contact_items:
                r = contact_p.add_run("  •  ")
                r.font.size = Pt(9)
                r.font.color.rgb = RGBColor(0x66, 0x66, 0x66)
            self._add_hyperlink(contact_p, self.info["website"], self.info["website"], font_size=9)
            contact_items.append(True)

        self._add_hr(doc)

        def add_section(title, items_func):
            heading = doc.add_paragraph()
            heading.paragraph_format.space_before = Pt(6)
            heading.paragraph_format.space_after = Pt(3)
            run = heading.add_run(title.upper())
            run.bold = True
            run.font.size = Pt(11)
            run.font.color.rgb = RGBColor(0x1A, 0x1A, 0x2E)
            self._add_hr(doc)
            items_func()

        # --- PROFESSIONAL SUMMARY ---
        add_section("Professional Summary", lambda: (
            doc.add_paragraph(self.info["summary"])
        ))

        # --- TECHNICAL SKILLS ---
        add_section("Technical Skills", lambda: (
            self._build_skills_section(doc)
        ))

        # --- LATEST PORTFOLIO (user's structure: portfolio first) ---
        add_section("Latest Portfolio", lambda: (
            self._build_portfolio_section(doc)
        ))

        # --- WORK EXPERIENCE ---
        add_section("Work Experience", lambda: (
            self._build_experience_section(doc)
        ))

        # --- PROJECTS ---
        if self.info.get("projects"):
            add_section("Projects", lambda: (
                self._build_projects_section(doc)
            ))

        # --- EDUCATION ---
        add_section("Education", lambda: (
            self._build_education_section(doc)
        ))

        doc.save(output_path)
        return output_path

    def _build_skills_section(self, doc):
        p = doc.add_paragraph()
        p.paragraph_format.space_before = Pt(1)
        skills_display = []
        for s in self.info["skills"]:
            skills_display.append(s)
        label_run = p.add_run("Languages & Frameworks: ")
        label_run.bold = True
        label_run.font.size = Pt(9.5)
        val_run = p.add_run(", ".join(skills_display[:8]))
        val_run.font.size = Pt(9.5)

        p2 = doc.add_paragraph()
        label2 = p2.add_run("Architecture & Patterns: ")
        label2.bold = True
        label2.font.size = Pt(9.5)
        val2 = p2.add_run(", ".join(skills_display[8:16]))
        val2.font.size = Pt(9.5)

        p3 = doc.add_paragraph()
        label3 = p3.add_run("Tools & Platforms: ")
        label3.bold = True
        label3.font.size = Pt(9.5)
        val3 = p3.add_run(", ".join(skills_display[16:]))
        val3.font.size = Pt(9.5)

    def _build_portfolio_section(self, doc):
        for i, project in enumerate(self.info.get("latest_portfolio", [])):
            # Title with links
            p = doc.add_paragraph()
            p.paragraph_format.space_before = Pt(5)
            p.paragraph_format.space_after = Pt(1)
            title_run = p.add_run(project["title"])
            title_run.bold = True
            title_run.font.size = Pt(10.5)
            title_run.font.color.rgb = RGBColor(0x1A, 0x1A, 0x2E)

            # Links inline
            if project.get("links"):
                link_p = doc.add_paragraph()
                link_p.paragraph_format.space_after = Pt(1)
                link_idx = 0
                for platform, url in project["links"].items():
                    if link_idx > 0:
                        sep = link_p.add_run("  |  ")
                        sep.font.size = Pt(8.5)
                        sep.font.color.rgb = RGBColor(0x66, 0x66, 0x66)
                    self._add_hyperlink(link_p, platform, url, font_size=8.5)
                    link_idx += 1

            # Description
            if project.get("description"):
                desc_p = doc.add_paragraph()
                desc_p.paragraph_format.space_after = Pt(1)
                desc_run = desc_p.add_run(project["description"])
                desc_run.italic = True
                desc_run.font.size = Pt(9.5)

            # Highlights
            for h in project.get("highlights", []):
                bp = doc.add_paragraph()
                bp.paragraph_format.space_before = Pt(0)
                bp.paragraph_format.space_after = Pt(1)
                bp.paragraph_format.left_indent = Inches(0.2)
                bullet = bp.add_run("\u2022  ")
                bullet.font.size = Pt(9)
                bullet.font.color.rgb = RGBColor(0x66, 0x66, 0x66)
                text_run = bp.add_run(h)
                text_run.font.size = Pt(9.5)

    def _build_experience_section(self, doc):
        for exp in self.info["experience"]:
            p = doc.add_paragraph()
            p.paragraph_format.space_before = Pt(5)
            p.paragraph_format.space_after = Pt(1)
            role_run = p.add_run(f"{exp['role']}")
            role_run.bold = True
            role_run.font.size = Pt(10.5)
            role_run.font.color.rgb = RGBColor(0x1A, 0x1A, 0x2E)
            sep = p.add_run("  |  ")
            sep.font.size = Pt(9)
            comp_run = p.add_run(exp["company"])
            comp_run.font.size = Pt(10)
            comp_run.font.color.rgb = RGBColor(0x55, 0x55, 0x55)

            p2 = doc.add_paragraph()
            p2.paragraph_format.space_after = Pt(2)
            per_run = p2.add_run(f"{exp['period']}  |  {exp['location']}")
            per_run.font.size = Pt(9)
            per_run.font.color.rgb = RGBColor(0x77, 0x77, 0x77)
            per_run.italic = True

            for h in exp["highlights"]:
                bp = doc.add_paragraph()
                bp.paragraph_format.space_before = Pt(0)
                bp.paragraph_format.space_after = Pt(1)
                bp.paragraph_format.left_indent = Inches(0.2)
                bullet = bp.add_run("\u2022  ")
                bullet.font.size = Pt(9)
                bullet.font.color.rgb = RGBColor(0x66, 0x66, 0x66)
                text_run = bp.add_run(h)
                text_run.font.size = Pt(9.5)

    def _build_projects_section(self, doc):
        for project in self.info.get("projects", []):
            p = doc.add_paragraph()
            p.paragraph_format.space_before = Pt(5)
            p.paragraph_format.space_after = Pt(1)
            title_run = p.add_run(project["title"])
            title_run.bold = True
            title_run.font.size = Pt(10.5)
            title_run.font.color.rgb = RGBColor(0x1A, 0x1A, 0x2E)
            if project.get("company"):
                sep_p = p.add_run("  -  ")
                comp_run = p.add_run(project["company"])
                comp_run.font.size = Pt(10)
                comp_run.font.color.rgb = RGBColor(0x55, 0x55, 0x55)

            if project.get("links"):
                link_p = doc.add_paragraph()
                link_p.paragraph_format.space_after = Pt(1)
                link_idx = 0
                for platform, url in project["links"].items():
                    if link_idx > 0:
                        sep = link_p.add_run("  |  ")
                        sep.font.size = Pt(8.5)
                        sep.font.color.rgb = RGBColor(0x66, 0x66, 0x66)
                    self._add_hyperlink(link_p, platform, url, font_size=8.5)
                    link_idx += 1

            for h in project.get("highlights", []):
                bp = doc.add_paragraph()
                bp.paragraph_format.space_before = Pt(0)
                bp.paragraph_format.space_after = Pt(1)
                bp.paragraph_format.left_indent = Inches(0.2)
                bullet = bp.add_run("\u2022  ")
                bullet.font.size = Pt(9)
                bullet.font.color.rgb = RGBColor(0x66, 0x66, 0x66)
                text_run = bp.add_run(h)
                text_run.font.size = Pt(9.5)

    def _build_education_section(self, doc):
        for edu in self.info["education"]:
            p = doc.add_paragraph()
            p.paragraph_format.space_before = Pt(3)
            deg_run = p.add_run(edu["degree"])
            deg_run.bold = True
            deg_run.font.size = Pt(10.5)
            if edu.get("school"):
                sep = p.add_run("  \u2014  ")
                sch_run = p.add_run(edu["school"])
                sch_run.font.size = Pt(10)
                sch_run.font.color.rgb = RGBColor(0x55, 0x55, 0x55)

            if edu.get("period") or edu.get("gpa"):
                p2 = doc.add_paragraph()
                p2.paragraph_format.space_after = Pt(1)
                parts = [x for x in [edu.get("period", ""), edu.get("gpa", "")] if x]
                if parts:
                    detail_run = p2.add_run("  |  ".join(parts))
                    detail_run.font.size = Pt(9)
                    detail_run.font.color.rgb = RGBColor(0x77, 0x77, 0x77)
                    detail_run.italic = True

    def build_pdf(self, output_path: str) -> str:
        class CVPDF(FPDF):
            def header(self):
                pass
            def footer(self):
                self.set_y(-15)
                self.set_font("Helvetica", "I", 7)
                self.set_text_color(150, 150, 150)
                self.cell(0, 10, f"Page {self.page_no()}/{{nb}}", align="C")

        pdf = CVPDF()
        pdf.alias_nb_pages()
        pdf.set_auto_page_break(auto=True, margin=20)
        pdf.add_page()
        pdf.set_margins(15, 12, 15)

        # --- HEADER ---
        pdf.set_font("Helvetica", "B", 22)
        pdf.set_text_color(26, 26, 46)
        pdf.cell(0, 9, self.info["name"], align="C", new_x="LMARGIN", new_y="NEXT")
        pdf.set_font("Helvetica", "", 11)
        pdf.set_text_color(85, 85, 85)
        pdf.cell(0, 7, self.info["title"], align="C", new_x="LMARGIN", new_y="NEXT")

        contact_parts = []
        if self.info.get("location"):
            contact_parts.append(self.info["location"])
        if self.info.get("linkedin"):
            contact_parts.append(f"LinkedIn")
        if self.info.get("github"):
            contact_parts.append("GitHub")
        if self.info.get("email"):
            contact_parts.append(self.info["email"])
        pdf.set_font("Helvetica", "", 8.5)
        pdf.set_text_color(102, 102, 102)
        pdf.cell(0, 6, "  |  ".join(contact_parts), align="C", new_x="LMARGIN", new_y="NEXT")
        pdf.ln(3)
        self._pdf_hr(pdf)
        pdf.ln(2)

        # --- SUMMARY ---
        self._pdf_section_header(pdf, "Professional Summary")
        pdf.set_font("Helvetica", "", 9.5)
        pdf.set_text_color(51, 51, 51)
        pdf.multi_cell(0, 5, self.info["summary"])
        pdf.ln(2)

        # --- SKILLS ---
        self._pdf_section_header(pdf, "Technical Skills")
        pdf.set_font("Helvetica", "", 9)
        pdf.set_text_color(51, 51, 51)
        skills_text = ", ".join(self.info["skills"])
        pdf.multi_cell(0, 4.5, skills_text)
        pdf.ln(2)

        # --- PORTFOLIO ---
        self._pdf_section_header(pdf, "Latest Portfolio")
        for project in self.info.get("latest_portfolio", []):
            pdf.set_font("Helvetica", "B", 10)
            pdf.set_text_color(26, 26, 46)
            title = project["title"]
            if project.get("links"):
                link_str = " | ".join(project["links"].keys())
                title += f"  ({link_str})"
            pdf.cell(0, 5.5, title, new_x="LMARGIN", new_y="NEXT")
            if project.get("description"):
                pdf.set_font("Helvetica", "I", 9)
                pdf.set_text_color(85, 85, 85)
                pdf.multi_cell(0, 4.5, project["description"])
            pdf.ln(1)
            pdf.set_font("Helvetica", "", 9)
            pdf.set_text_color(51, 51, 51)
            for h in project.get("highlights", []):
                x = pdf.get_x()
                pdf.cell(5, 4.5, "-")
                pdf.multi_cell(0, 4.5, f" {h}")
                pdf.set_x(x)
            pdf.ln(2)

        # --- WORK EXPERIENCE ---
        self._pdf_section_header(pdf, "Work Experience")
        for exp in self.info["experience"]:
            pdf.set_font("Helvetica", "B", 10)
            pdf.set_text_color(26, 26, 46)
            pdf.cell(0, 5.5, f"{exp['role']}  |  {exp['company']}", new_x="LMARGIN", new_y="NEXT")
            pdf.set_font("Helvetica", "I", 8.5)
            pdf.set_text_color(119, 119, 119)
            pdf.cell(0, 4.5, f"{exp['period']}  |  {exp['location']}", new_x="LMARGIN", new_y="NEXT")
            pdf.ln(1)
            pdf.set_font("Helvetica", "", 9)
            pdf.set_text_color(51, 51, 51)
            for h in exp["highlights"]:
                x = pdf.get_x()
                pdf.cell(5, 4.5, "-")
                pdf.multi_cell(0, 4.5, f" {h}")
                pdf.set_x(x)
            pdf.ln(2)

        # --- PROJECTS ---
        if self.info.get("projects"):
            self._pdf_section_header(pdf, "Projects")
            for proj in self.info["projects"]:
                pdf.set_font("Helvetica", "B", 10)
                pdf.set_text_color(26, 26, 46)
                proj_title = proj["title"]
                if proj.get("company"):
                    proj_title += f"  -  {proj['company']}"
                if proj.get("links"):
                    link_str = " | ".join(proj["links"].keys())
                    proj_title += f"  ({link_str})"
                pdf.cell(0, 5.5, proj_title, new_x="LMARGIN", new_y="NEXT")
                pdf.ln(1)
                pdf.set_font("Helvetica", "", 9)
                pdf.set_text_color(51, 51, 51)
                for h in proj.get("highlights", []):
                    x = pdf.get_x()
                    pdf.cell(5, 4.5, "-")
                    pdf.multi_cell(0, 4.5, f" {h}")
                    pdf.set_x(x)
                pdf.ln(2)

        # --- EDUCATION ---
        self._pdf_section_header(pdf, "Education")
        for edu in self.info["education"]:
            pdf.set_font("Helvetica", "B", 10)
            pdf.set_text_color(26, 26, 46)
            edu_line = edu["degree"]
            if edu.get("school"):
                edu_line += f"  -  {edu['school']}"
            pdf.cell(0, 5.5, edu_line, new_x="LMARGIN", new_y="NEXT")
            if edu.get("period") or edu.get("gpa"):
                pdf.set_font("Helvetica", "I", 8.5)
                pdf.set_text_color(119, 119, 119)
                parts = [x for x in [edu.get("period", ""), edu.get("gpa", "")] if x]
                if parts:
                    pdf.cell(0, 4.5, "  |  ".join(parts), new_x="LMARGIN", new_y="NEXT")
            pdf.ln(2)

        pdf.output(output_path)
        return output_path

    def _pdf_section_header(self, pdf, title):
        pdf.set_font("Helvetica", "B", 11)
        pdf.set_text_color(26, 26, 46)
        pdf.cell(0, 7, title.upper(), new_x="LMARGIN", new_y="NEXT")
        self._pdf_hr(pdf)
        pdf.ln(2)

    def _pdf_hr(self, pdf):
        y = pdf.get_y()
        pdf.set_draw_color(51, 51, 51)
        pdf.set_line_width(0.3)
        pdf.line(15, y, pdf.w - 15, y)
        pdf.set_y(y + 1.5)

    def get_keyword_density(self) -> dict:
        text = self.info["summary"] + " "
        text += " ".join(self.info["skills"]) + " "
        text += " ".join(
            h for exp in self.info["experience"] for h in exp["highlights"]
        )
        text += " ".join(
            h for p in self.info.get("latest_portfolio", [])
            for h in p.get("highlights", [])
        )
        text += " ".join(
            h for p in self.info.get("projects", [])
            for h in p.get("highlights", [])
        )
        result = {}
        for kw in self.ATS_KEYWORDS:
            count = len(re.findall(re.escape(kw), text, re.IGNORECASE))
            result[kw] = count
        return result
