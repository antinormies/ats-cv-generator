from docx import Document
from docx.shared import Pt, Inches, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml.ns import qn
from docx.oxml import OxmlElement
from fpdf import FPDF
from datetime import date
import os


class CoverLetterBuilder:
    def __init__(self, info: dict):
        self.info = info

    def _sanitize_text(self, text: str) -> str:
        replacements = {
            "\u2019": "'", "\u2018": "'",
            "\u201C": '"', "\u201D": '"',
            "\u2013": "-", "\u2014": "-",
            "\u2022": "-", "\u2026": "...",
            "\u00A0": " ",
        }
        for old, new in replacements.items():
            text = text.replace(old, new)
        return text

    def _normalize_url(self, url: str) -> str:
        if url.startswith("mailto:"):
            return url
        if not url.startswith(("http://", "https://")):
            return f"https://{url}"
        return url

    def _add_hyperlink(self, paragraph, text: str, url: str, font_size=9):
        url = self._normalize_url(url)
        part = paragraph.part
        r_id = part.relate_to(url, "http://schemas.openxmlformats.org/officeDocument/2006/relationships/hyperlink", is_external=True)
        hyperlink = OxmlElement("w:hyperlink")
        hyperlink.set(qn("r:id"), r_id)
        new_run = OxmlElement("w:r")
        rPr = OxmlElement("w:rPr")
        c = OxmlElement("w:color")
        c.set(qn("w:val"), "2255CC")
        rPr.append(c)
        u = OxmlElement("w:u")
        u.set(qn("w:val"), "single")
        rPr.append(u)
        sz = OxmlElement("w:sz")
        sz.set(qn("w:val"), str(int(font_size * 2)))
        rPr.append(sz)
        new_run.append(rPr)
        t = OxmlElement("w:t")
        t.text = text
        new_run.append(t)
        hyperlink.append(new_run)
        paragraph._p.append(hyperlink)
        return hyperlink

    def _add_hr(self, doc):
        p = doc.add_paragraph()
        p.paragraph_format.space_before = Pt(2)
        p.paragraph_format.space_after = Pt(2)
        pPr = p._p.get_or_add_pPr()
        pBdr = OxmlElement("w:pBdr")
        bottom = OxmlElement("w:bottom")
        bottom.set(qn("w:val"), "single")
        bottom.set(qn("w:sz"), "4")
        bottom.set(qn("w:space"), "1")
        bottom.set(qn("w:color"), "333333")
        pBdr.append(bottom)
        pPr.append(pBdr)

    def build(self, output_path: str, body: str, company: str = "", recipient: str = "") -> str:
        doc = Document()
        section = doc.sections[0]
        section.top_margin = Inches(0.5)
        section.bottom_margin = Inches(0.5)
        section.left_margin = Inches(0.7)
        section.right_margin = Inches(0.7)

        style = doc.styles["Normal"]
        font = style.font
        font.name = "Calibri"
        font.size = Pt(10.5)
        font.color.rgb = RGBColor(0x33, 0x33, 0x33)
        style.paragraph_format.space_after = Pt(3)
        style.paragraph_format.space_before = Pt(0)
        style.paragraph_format.line_spacing = 1.08

        # --- HEADER (same as CV) ---
        name_p = doc.add_paragraph()
        name_p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        name_run = name_p.add_run(self.info["name"])
        name_run.bold = True
        name_run.font.size = Pt(20)
        name_run.font.color.rgb = RGBColor(0x1A, 0x1A, 0x2E)

        title_p = doc.add_paragraph()
        title_p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        title_run = title_p.add_run(self.info["title"])
        title_run.font.size = Pt(11)
        title_run.font.color.rgb = RGBColor(0x55, 0x55, 0x55)

        contact_p = doc.add_paragraph()
        contact_p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        contact_items = []
        if self.info["location"]:
            r = contact_p.add_run(self.info["location"])
            r.font.size = Pt(9)
            r.font.color.rgb = RGBColor(0x66, 0x66, 0x66)
            contact_items.append(True)
        if self.info["linkedin"]:
            if contact_items:
                r = contact_p.add_run("  \u2022  ")
                r.font.size = Pt(9)
                r.font.color.rgb = RGBColor(0x66, 0x66, 0x66)
            self._add_hyperlink(contact_p, "LinkedIn", self.info["linkedin"], font_size=9)
            contact_items.append(True)
        if self.info["github"]:
            if contact_items:
                r = contact_p.add_run("  \u2022  ")
                r.font.size = Pt(9)
                r.font.color.rgb = RGBColor(0x66, 0x66, 0x66)
            self._add_hyperlink(contact_p, "GitHub", self.info["github"], font_size=9)
            contact_items.append(True)
        if self.info["email"]:
            if contact_items:
                r = contact_p.add_run("  \u2022  ")
                r.font.size = Pt(9)
                r.font.color.rgb = RGBColor(0x66, 0x66, 0x66)
            self._add_hyperlink(contact_p, self.info["email"], f"mailto:{self.info['email']}", font_size=9)
            contact_items.append(True)
        if self.info["website"]:
            if contact_items:
                r = contact_p.add_run("  \u2022  ")
                r.font.size = Pt(9)
                r.font.color.rgb = RGBColor(0x66, 0x66, 0x66)
            self._add_hyperlink(contact_p, self.info["website"], self.info["website"], font_size=9)
            contact_items.append(True)

        self._add_hr(doc)

        # --- DATE ---
        from datetime import date
        date_p = doc.add_paragraph()
        date_p.paragraph_format.space_before = Pt(8)
        date_p.paragraph_format.space_after = Pt(6)
        date_run = date_p.add_run(date.today().strftime("%B %d, %Y"))
        date_run.font.size = Pt(10)

        # --- RECIPIENT ---
        if recipient:
            rp = doc.add_paragraph()
            rp.paragraph_format.space_after = Pt(1)
            rr = rp.add_run(recipient)
            rr.font.size = Pt(10)
        if company:
            cp = doc.add_paragraph()
            cp.paragraph_format.space_after = Pt(6)
            cr = cp.add_run(company)
            cr.font.size = Pt(10)

        # --- SUBJECT ---
        subj_p = doc.add_paragraph()
        subj_p.paragraph_format.space_after = Pt(6)
        subj_r = subj_p.add_run(f"Re: Application for {self.info['title']} Position")
        subj_r.bold = True
        subj_r.font.size = Pt(10.5)

        # --- SALUTATION ---
        sal_p = doc.add_paragraph()
        sal_p.paragraph_format.space_after = Pt(6)
        sal_r = sal_p.add_run("Dear Hiring Manager,")
        sal_r.font.size = Pt(10.5)

        # --- BODY ---
        paragraphs = [p.strip() for p in body.split("\n") if p.strip()]
        for para_text in paragraphs:
            bp = doc.add_paragraph()
            bp.paragraph_format.space_after = Pt(6)
            bp.paragraph_format.line_spacing = 1.15
            br = bp.add_run(para_text)
            br.font.size = Pt(10.5)

        # --- CLOSING ---
        close_p = doc.add_paragraph()
        close_p.paragraph_format.space_before = Pt(6)
        close_p.paragraph_format.space_after = Pt(2)
        close_r = close_p.add_run("Best regards,")
        close_r.font.size = Pt(10.5)

        name_close = doc.add_paragraph()
        name_close.paragraph_format.space_after = Pt(1)
        ncr = name_close.add_run(self.info["name"])
        ncr.bold = True
        ncr.font.size = Pt(10.5)

        if self.info.get("phone"):
            phone_p = doc.add_paragraph()
            phone_p.paragraph_format.space_after = Pt(1)
            ph = phone_p.add_run(self.info["phone"])
            ph.font.size = Pt(9.5)
            ph.font.color.rgb = RGBColor(0x66, 0x66, 0x66)

        if self.info.get("email"):
            email_p = doc.add_paragraph()
            email_p.paragraph_format.space_after = Pt(1)
            em = email_p.add_run(self.info["email"])
            em.font.size = Pt(9.5)
            em.font.color.rgb = RGBColor(0x66, 0x66, 0x66)

        doc.save(output_path)
        return output_path

    def build_pdf(self, output_path: str, body: str, company: str = "", recipient: str = "") -> str:
        class CLPDF(FPDF):
            def header(self):
                pass
            def footer(self):
                self.set_y(-15)
                self.set_font("Helvetica", "I", 7)
                self.set_text_color(150, 150, 150)
                self.cell(0, 10, f"Page {self.page_no()}/{{nb}}", align="C")

        pdf = CLPDF()
        pdf.alias_nb_pages()
        pdf.set_auto_page_break(auto=True, margin=20)
        pdf.add_page()
        pdf.set_margins(15, 12, 15)

        # --- HEADER ---
        pdf.set_font("Helvetica", "B", 22)
        pdf.set_text_color(26, 26, 46)
        pdf.cell(0, 9, self.info["name"], align="C", new_x="LMARGIN", new_y="NEXT")
        pdf.set_font("Helvetica", "", 11)
        pdf.set_text_color(85, 85, 85)
        pdf.cell(0, 7, self.info["title"], align="C", new_x="LMARGIN", new_y="NEXT")

        pdf.set_font("Helvetica", "", 8.5)
        pdf.set_text_color(102, 102, 102)
        contact_items = []
        if self.info.get("location"):
            contact_items.append(("text", self.info["location"], None))
        if self.info.get("linkedin"):
            contact_items.append(("link", "LinkedIn", self._normalize_url(self.info["linkedin"])))
        if self.info.get("github"):
            contact_items.append(("link", "GitHub", self._normalize_url(self.info["github"])))
        if self.info.get("email"):
            contact_items.append(("link", self.info["email"], f"mailto:{self.info['email']}"))
        if self.info.get("website"):
            contact_items.append(("link", self.info["website"], self._normalize_url(self.info["website"])))
        total_w = 0
        for idx, (ctype, text, url) in enumerate(contact_items):
            if idx > 0:
                total_w += pdf.get_string_width("  |  ")
            total_w += pdf.get_string_width(text)
        x_cur = (pdf.w - total_w) / 2
        y_pos = pdf.get_y()
        for idx, (ctype, text, url) in enumerate(contact_items):
            if idx > 0:
                sep = "  |  "
                sw = pdf.get_string_width(sep)
                pdf.set_xy(x_cur, y_pos)
                pdf.cell(sw, 6, sep)
                x_cur += sw
            tw = pdf.get_string_width(text)
            pdf.set_xy(x_cur, y_pos)
            if ctype == "link":
                pdf.set_text_color(34, 85, 204)
                pdf.cell(tw, 6, text)
                pdf.link(x_cur, y_pos, tw, 6, url)
                pdf.set_text_color(102, 102, 102)
            else:
                pdf.cell(tw, 6, text)
            x_cur += tw
        pdf.set_y(y_pos + 8)

        # --- HR ---
        y = pdf.get_y()
        pdf.set_draw_color(51, 51, 51)
        pdf.set_line_width(0.3)
        pdf.line(15, y, pdf.w - 15, y)
        pdf.set_y(y + 4)

        # --- DATE ---
        pdf.set_font("Helvetica", "", 10)
        pdf.set_text_color(51, 51, 51)
        pdf.cell(0, 6, self._sanitize_text(date.today().strftime("%B %d, %Y")), new_x="LMARGIN", new_y="NEXT")
        pdf.ln(2)

        # --- RECIPIENT ---
        if recipient:
            pdf.cell(0, 5.5, recipient, new_x="LMARGIN", new_y="NEXT")
        if company:
            pdf.cell(0, 5.5, company, new_x="LMARGIN", new_y="NEXT")

        pdf.ln(2)

        # --- SUBJECT ---
        pdf.set_font("Helvetica", "B", 10.5)
        pdf.cell(0, 6, self._sanitize_text(f"Re: Application for {self.info['title']} Position"), new_x="LMARGIN", new_y="NEXT")
        pdf.ln(3)

        # --- SALUTATION ---
        pdf.set_font("Helvetica", "", 10.5)
        pdf.cell(0, 6, "Dear Hiring Manager,", new_x="LMARGIN", new_y="NEXT")
        pdf.ln(3)

        # --- BODY ---
        pdf.set_font("Helvetica", "", 10.5)
        for para in body.split("\n"):
            para = self._sanitize_text(para.strip())
            if para:
                pdf.multi_cell(0, 5.5, para)
                pdf.ln(2)

        # --- CLOSING ---
        pdf.ln(4)
        pdf.cell(0, 6, "Best regards,", new_x="LMARGIN", new_y="NEXT")
        pdf.ln(2)
        pdf.set_font("Helvetica", "B", 10.5)
        pdf.cell(0, 6, self.info["name"], new_x="LMARGIN", new_y="NEXT")
        pdf.set_font("Helvetica", "", 9.5)
        pdf.set_text_color(102, 102, 102)
        if self.info.get("phone"):
            pdf.cell(0, 5, self.info["phone"], new_x="LMARGIN", new_y="NEXT")
        if self.info.get("email"):
            pdf.cell(0, 5, self.info["email"], new_x="LMARGIN", new_y="NEXT")

        pdf.output(output_path)
        return output_path
